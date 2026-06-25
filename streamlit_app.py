"""
Doküman Asistanı — Streamlit Cloud sürümü (tek dosya, kendi içinde tam).

Ayarları Streamlit Cloud Secrets'tan okur. Yerel sürümden farkları:
- logic/ paketine bağımlı DEĞİL (import yolu sorunu olmaz).
- ChromaDB KULLANMAZ — bunun yerine hafif bellek-içi vektör deposu kullanır
  (Streamlit Cloud'da Python 3.14 + protobuf ile chromadb çöküyor; ayrıca
  Cloud diski geçici olduğu için kalıcı DB'nin de anlamı yok).

NOT: Yüklenen dokümanlar yalnızca o oturumda hafızada tutulur; uygulama
yeniden başlarsa tekrar yüklemen gerekir.
"""

import io
import re
from dataclasses import dataclass

import numpy as np
import streamlit as st

st.set_page_config(page_title="Doküman Asistanı", layout="wide")


# === Ayarlar (Secrets) ====================================================
def _secret(key: str, default: str = "") -> str:
    try:
        value = st.secrets.get(key, default)
    except Exception:
        value = default
    return str(value).strip() if value else default


LLM_BASE_URL = _secret("LLM_BASE_URL")
LLM_MODEL = _secret("LLM_MODEL")
LLM_API_KEY = _secret("LLM_API_KEY")
EMBEDDING_MODEL = _secret("EMBEDDING_MODEL", "intfloat/multilingual-e5-small")
LLM_TEMPERATURE = 0.2
TOP_K = 5
CANDIDATE_K = 30
CHUNK_SIZE = 1200
CHUNK_OVERLAP = 150
SUPPORTED_EXTENSIONS = (".pdf", ".docx", ".txt", ".md", ".html", ".htm")


# === Doküman okuma ========================================================
@dataclass
class Chunk:
    text: str
    document_name: str
    page_number: int


def _read_pdf(buf):
    from pypdf import PdfReader

    pages = []
    for i, page in enumerate(PdfReader(buf).pages, start=1):
        text = page.extract_text() or ""
        if text.strip():
            pages.append((text, i))
    return pages


def _read_docx(buf):
    from docx import Document

    doc = Document(buf)
    parts = [p.text for p in doc.paragraphs if p.text.strip()]
    for table in doc.tables:
        for row in table.rows:
            cells = [c.text.strip() for c in row.cells if c.text.strip()]
            if cells:
                parts.append(" | ".join(cells))
    text = "\n\n".join(parts)
    return [(text, 1)] if text else []


def _read_txt(buf):
    raw = buf.read()
    try:
        text = raw.decode("utf-8")
    except UnicodeDecodeError:
        text = raw.decode("latin-1", errors="replace")
    return [(text, 1)] if text.strip() else []


def _read_markdown(buf):
    from bs4 import BeautifulSoup
    from markdown import markdown

    html = markdown(buf.read().decode("utf-8", errors="replace"))
    text = BeautifulSoup(html, "html.parser").get_text(separator="\n")
    return [(text, 1)] if text.strip() else []


def _read_html(buf):
    from bs4 import BeautifulSoup

    soup = BeautifulSoup(buf.read(), "html.parser")
    for tag in soup(["script", "style", "noscript"]):
        tag.decompose()
    text = "\n".join(line.strip() for line in soup.get_text("\n").splitlines() if line.strip())
    return [(text, 1)] if text else []


_READERS = {
    ".pdf": _read_pdf, ".docx": _read_docx, ".txt": _read_txt,
    ".md": _read_markdown, ".html": _read_html, ".htm": _read_html,
}


def load_document(name, data):
    name_l = name.lower()
    for ext, reader in _READERS.items():
        if name_l.endswith(ext):
            return reader(io.BytesIO(data))
    raise ValueError(f"Desteklenmeyen format: {name}")


# Mevzuat madde başlıkları: "MADDE 5", "GEÇİCİ MADDE 1", "EK MADDE 2".
_ARTICLE_RE = re.compile(
    r"(?=^\s*(?:GEÇİCİ\s+|EK\s+)?MADDE\s+\d+)", re.IGNORECASE | re.MULTILINE
)
# Cümle sonu: nokta/soru/ünlem + boşluk + büyük harf/rakam ile başlayan yeni cümle.
_SENTENCE_RE = re.compile(r"""(?<=[.!?…])["'”’\)\]]*\s+(?=[A-ZÇĞİÖŞÜ0-9])""")


def _split_sentences(text):
    sentences = []
    for line in text.split("\n"):
        line = line.strip()
        if not line:
            continue
        for sent in _SENTENCE_RE.split(line):
            sent = sent.strip()
            if sent:
                sentences.append(sent)
    return sentences


def _hard_split(sentence, max_chars):
    parts, start = [], 0
    while start < len(sentence):
        end = min(len(sentence), start + max_chars)
        if end < len(sentence):
            cut = sentence.rfind(" ", start + max_chars // 2, end)
            if cut != -1:
                end = cut
        parts.append(sentence[start:end].strip())
        start = end
    return [p for p in parts if p]


def _pack(sentences, max_chars, overlap):
    chunks, current, current_len = [], [], 0
    for sent in sentences:
        if len(sent) > max_chars:
            if current:
                chunks.append(" ".join(current))
                current, current_len = [], 0
            chunks.extend(_hard_split(sent, max_chars))
            continue
        addition = len(sent) + (1 if current else 0)
        if current and current_len + addition > max_chars:
            chunks.append(" ".join(current))
            carry, carry_len = [], 0
            for prev in reversed(current):
                if carry_len + len(prev) > overlap and carry:
                    break
                carry.insert(0, prev)
                carry_len += len(prev) + 1
            current = carry
            current_len = sum(len(s) + 1 for s in current)
        current.append(sent)
        current_len += addition
    if current:
        chunks.append(" ".join(current))
    return [c.strip() for c in chunks if c.strip()]


def chunk_text(text, max_chars=CHUNK_SIZE, overlap=CHUNK_OVERLAP):
    """Madde-farkındalıklı, cümle-bütünlüğü korunan parçalama (yarım cümle bırakmaz)."""
    if len(text) <= max_chars:
        stripped = text.strip()
        return [stripped] if stripped else []
    blocks = [b for b in _ARTICLE_RE.split(text) if b.strip()] or [text]
    parts = []
    for block in blocks:
        sentences = _split_sentences(block)
        if sentences:
            parts.extend(_pack(sentences, max_chars, overlap))
    return parts


# === Embedding (lokal, ücretsiz) ==========================================
@st.cache_resource(show_spinner=False)
def get_embedder(model_name):
    from sentence_transformers import SentenceTransformer

    return SentenceTransformer(model_name)


def embed_texts(texts, prefix):
    model = get_embedder(EMBEDDING_MODEL)
    vecs = model.encode([f"{prefix}: {t}" for t in texts],
                        normalize_embeddings=True, show_progress_bar=False)
    return np.asarray(vecs, dtype=np.float32)


# === Bellek-içi arama (vektör + BM25 hybrid) ==============================
def _tokenize(text):
    return re.findall(r"\w+", text.lower())


def retrieve(question, chunks, vectors):
    if not chunks:
        return []
    q_vec = embed_texts([question], "query")[0]
    sims = vectors @ q_vec  # normalize edilmiş → kosinüs benzerliği
    vec_rank = np.argsort(-sims)[:CANDIDATE_K]

    from rank_bm25 import BM25Okapi

    bm25 = BM25Okapi([_tokenize(c.text) for c in chunks])
    bm_scores = bm25.get_scores(_tokenize(question))
    bm_rank = np.argsort(-bm_scores)[:CANDIDATE_K]

    # Reciprocal Rank Fusion
    scores = {}
    for rank, idx in enumerate(vec_rank):
        scores[idx] = scores.get(idx, 0.0) + 1.0 / (60 + rank + 1)
    for rank, idx in enumerate(bm_rank):
        scores[idx] = scores.get(idx, 0.0) + 1.0 / (60 + rank + 1)
    best = sorted(scores, key=lambda i: scores[i], reverse=True)[:TOP_K]
    return [chunks[i] for i in best]


# === LLM (st.secrets + openai SDK — AI Hub streamlit kalıbı) ==============
@st.cache_resource(show_spinner=False)
def _llm_client(base_url, api_key):
    from openai import OpenAI

    # api_key boş olsa bile (anahtarsız lokal uç noktalar için) istemci kurulur.
    return OpenAI(api_key=api_key or "no-key", base_url=base_url or None)


def llm_chat(messages):
    client = _llm_client(LLM_BASE_URL, LLM_API_KEY)
    resp = client.chat.completions.create(
        model=LLM_MODEL, messages=messages, temperature=LLM_TEMPERATURE,
    )
    return resp.choices[0].message.content


SYSTEM_PROMPT = (
    "Sen verilen dokümanlar üzerinden soru cevaplayan bir asistansın. "
    "Cevabını yalnızca sağlanan PASAJ'lara dayandır. "
    "Pasajlarda cevap yoksa açıkça 'Sağlanan dokümanlarda bu bilgi yok.' de. "
    "Kaynak pasaj numarasını köşeli parantezle belirt: [1], [2]. "
    "Türkçe sorulara Türkçe, İngilizce sorulara İngilizce cevap ver."
)


def answer(question, history, chunks, vectors):
    hits = retrieve(question, chunks, vectors)
    if not hits:
        return "Soruyla eşleşen pasaj bulunamadı.", []
    context = "\n\n".join(
        f"[{i}] (Kaynak: {c.document_name}, sayfa {c.page_number})\n{c.text}"
        for i, c in enumerate(hits, start=1)
    )
    messages = [{"role": "system", "content": SYSTEM_PROMPT}, *history,
                {"role": "user", "content": f"PASAJLAR:\n{context}\n\nSORU: {question}"}]
    return llm_chat(messages), hits


# === Arayüz ===============================================================
st.title("Doküman Asistanı")
st.caption("Dokümanlarını yükle, içerikleri hakkında günlük dille soru sor. "
           "Yanıtlar yalnızca senin yüklediğin belgelere dayanır.")

if not LLM_BASE_URL or not LLM_MODEL:
    st.error("Model bilgileri Secrets'tan okunamadı. Streamlit Cloud → **Manage app → "
             "Settings → Secrets** altına en az `LLM_BASE_URL` ve `LLM_MODEL` gir, "
             "sonra **Reboot** et.")
    with st.expander("Secrets durumu (değerler gizli)"):
        for k in ("LLM_BASE_URL", "LLM_MODEL", "LLM_API_KEY", "EMBEDDING_MODEL"):
            st.write(f"{'✓' if _secret(k) else '✗'} `{k}`")
    st.stop()

for key, default in (("messages", []), ("chunks", []), ("vectors", None), ("sources", {})):
    if key not in st.session_state:
        st.session_state[key] = default


def _preview(text, limit=800):
    if len(text) <= limit:
        return text
    head = text[:limit]
    cut = max(head.rfind(". "), head.rfind("\n"), head.rfind(" "))
    if cut > limit // 2:
        head = head[: cut + 1]
    return head.rstrip() + " …"


def render_sources(sources):
    if not sources:
        return
    with st.expander(f"Yanıtın dayandığı doküman bölümleri ({len(sources)})"):
        for c in sources:
            st.markdown(f"**{c.document_name} — sayfa {c.page_number}**")
            st.caption(_preview(c.text))


with st.sidebar:
    st.header("Doküman yükle")
    st.caption("Yüklenenler geçicidir; uygulama yeniden başlarsa tekrar yükle.")
    uploaded = st.file_uploader(
        "PDF, Word, metin (TXT), Markdown veya web sayfası (HTML)",
        type=[e.lstrip(".") for e in SUPPORTED_EXTENSIONS], accept_multiple_files=True,
    )
    if uploaded and st.button("Dokümanları ekle", type="primary", use_container_width=True):
        progress = st.progress(0.0)
        new_chunks = []
        for i, f in enumerate(uploaded):
            try:
                pages = load_document(f.name, f.read())
                for text, page in pages:
                    for part in chunk_text(text):
                        new_chunks.append(Chunk(part, f.name, page))
                st.success(f"'{f.name}' eklendi.")
            except Exception as e:
                st.error(f"'{f.name}' eklenirken sorun: {e}")
            progress.progress((i + 1) / len(uploaded))
        if new_chunks:
            with st.spinner("Belgeler işleniyor…"):
                new_vecs = embed_texts([c.text for c in new_chunks], "passage")
            st.session_state.chunks.extend(new_chunks)
            st.session_state.vectors = (
                new_vecs if st.session_state.vectors is None
                else np.vstack([st.session_state.vectors, new_vecs])
            )

    st.divider()
    st.header("Yüklenen dokümanlar")
    names = sorted({c.document_name for c in st.session_state.chunks})
    if not names:
        st.info("Henüz doküman yüklenmedi.")
    else:
        for n in names:
            st.write(f"• {n}")
        if st.button("Tümünü temizle", use_container_width=True):
            st.session_state.update(messages=[], chunks=[], vectors=None, sources={})
            st.rerun()


for i, msg in enumerate(st.session_state.messages):
    with st.chat_message(msg["role"]):
        st.markdown(msg["content"])
        render_sources(st.session_state.sources.get(i))


if prompt := st.chat_input("Dokümanların hakkında bir soru yaz…"):
    if not st.session_state.chunks:
        st.warning("Önce sol taraftan en az bir doküman yüklemelisin.")
        st.stop()
    st.session_state.messages.append({"role": "user", "content": prompt})
    with st.chat_message("user"):
        st.markdown(prompt)
    with st.chat_message("assistant"):
        with st.spinner("Yanıt hazırlanıyor…"):
            try:
                history = [{"role": m["role"], "content": m["content"]}
                           for m in st.session_state.messages[:-1]]
                text, sources = answer(prompt, history,
                                       st.session_state.chunks, st.session_state.vectors)
            except Exception as e:
                st.error(f"Yanıt alınamadı: {e}")
                st.stop()
        st.markdown(text)
        idx = len(st.session_state.messages)
        st.session_state.messages.append({"role": "assistant", "content": text})
        st.session_state.sources[idx] = sources
        render_sources(sources)
