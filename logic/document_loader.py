from __future__ import annotations

import io
from dataclasses import dataclass

from bs4 import BeautifulSoup
from docx import Document as DocxDocument
from markdown import markdown
from pypdf import PdfReader


SUPPORTED_EXTENSIONS = (".pdf", ".docx", ".txt", ".md", ".html", ".htm")


@dataclass
class DocumentPage:
    """Bir dokümandan çıkarılan tek bir mantıksal sayfa/blok."""

    text: str
    page_number: int  # 1-based; sayfa kavramı yoksa 1


def _read_pdf(buffer: io.BytesIO) -> list[DocumentPage]:
    reader = PdfReader(buffer)
    pages: list[DocumentPage] = []
    for i, page in enumerate(reader.pages, start=1):
        text = page.extract_text() or ""
        if text.strip():
            pages.append(DocumentPage(text=text, page_number=i))
    return pages


def _read_docx(buffer: io.BytesIO) -> list[DocumentPage]:
    doc = DocxDocument(buffer)
    paragraphs = [p.text for p in doc.paragraphs if p.text.strip()]
    for table in doc.tables:
        for row in table.rows:
            cells = [c.text.strip() for c in row.cells if c.text.strip()]
            if cells:
                paragraphs.append(" | ".join(cells))
    text = "\n\n".join(paragraphs)
    return [DocumentPage(text=text, page_number=1)] if text else []


def _read_txt(buffer: io.BytesIO) -> list[DocumentPage]:
    raw = buffer.read()
    try:
        text = raw.decode("utf-8")
    except UnicodeDecodeError:
        text = raw.decode("latin-1", errors="replace")
    return [DocumentPage(text=text, page_number=1)] if text.strip() else []


def _read_markdown(buffer: io.BytesIO) -> list[DocumentPage]:
    raw = buffer.read().decode("utf-8", errors="replace")
    html = markdown(raw)
    text = BeautifulSoup(html, "html.parser").get_text(separator="\n")
    return [DocumentPage(text=text, page_number=1)] if text.strip() else []


def _read_html(buffer: io.BytesIO) -> list[DocumentPage]:
    soup = BeautifulSoup(buffer.read(), "html.parser")
    for tag in soup(["script", "style", "noscript"]):
        tag.decompose()
    text = soup.get_text(separator="\n")
    text = "\n".join(line.strip() for line in text.splitlines() if line.strip())
    return [DocumentPage(text=text, page_number=1)] if text else []


_READERS = {
    ".pdf": _read_pdf,
    ".docx": _read_docx,
    ".txt": _read_txt,
    ".md": _read_markdown,
    ".html": _read_html,
    ".htm": _read_html,
}


def load_document(name: str, buffer) -> list[DocumentPage]:
    """Yüklenen dosyayı uzantıya göre doğru parser'a yönlendirir."""
    name_l = name.lower()
    for ext, reader in _READERS.items():
        if name_l.endswith(ext):
            data = buffer.read() if hasattr(buffer, "read") else buffer
            return reader(io.BytesIO(data))
    raise ValueError(
        f"Desteklenmeyen format: {name}. Desteklenenler: {', '.join(SUPPORTED_EXTENSIONS)}"
    )
